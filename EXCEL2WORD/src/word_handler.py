from docx import Document
from docx.shared import RGBColor
import os
from .utils import format_filename
from tqdm import tqdm

class WordHandler:
    def __init__(self, config):
        self.config = config
        # 使用配置中的excel_handler
        if not config.excel_handler:
            from .excel_handler import ExcelHandler
            config.excel_handler = ExcelHandler(config)
            if not config.excel_path:
                raise ValueError("Excel文件路径未设置")

    def process_documents(self):
        """处理所有文档"""
        try:
            # 读取Excel数据
            print("正在读取Excel数据...")
            df = self.config.excel_handler.read_data()
            total_rows = len(df)
            print(f"成功读取数据，共 {total_rows} 行")
            
            # 使用tqdm创建进度条
            with tqdm(total=total_rows, desc="处理进度", ncols=100) as pbar:
                # 处理每一行数据
                for idx, row in df.iterrows():
                    try:
                        self._process_single_document(row)
                        pbar.update(1)  # 更新进度条
                    except Exception as e:
                        print(f"\n处理第 {idx + 1} 行数据时出错: {str(e)}")
                        
        except Exception as e:
            raise Exception(f"处理文档失败: {str(e)}")

    def _process_single_document(self, row):
        """处理单个文档"""
        try:
            # 检查数据行但不打印详细信息
            for item in self.config.replace_items:
                if item not in row:
                    raise ValueError(f"数据中缺少列: {item}")

            # 读取模板
            doc = Document(self.config.word_template)
            
            # 替换文本
            self._replace_in_document(doc, row)
            
            # 生成输出文件名
            replace_dict = {item: row[item] for item in self.config.replace_items}
            output_filename = format_filename(self.config.output_format, replace_dict)
            output_path = os.path.join(self.config.output_dir, f"{output_filename}.docx")
            
            # 保存文档，不打印详细信息
            doc.save(output_path)
            
        except Exception as e:
            raise Exception(f"处理文档失败: {str(e)}")

    def _replace_in_document(self, doc, row):
        """在文档中替换文本"""
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, row)
            
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, row)

    def _replace_in_paragraph(self, paragraph, data_row):
        """在段落中替换文本"""
        for item in self.config.replace_items:
            if item in paragraph.text:
                value = str(data_row[item])
                
                # 保存原始格式
                runs = paragraph.runs
                for run in runs:
                    if item in run.text:
                        run.text = run.text.replace(item, value)
                        run.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色 