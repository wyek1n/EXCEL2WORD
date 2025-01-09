import os
import platform
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import red, black
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

class PDFConverter:
    def __init__(self, config):
        self.config = config
        self.font_name = 'CustomFont'
        self.setup_fonts()

    def setup_fonts(self):
        """设置字体"""
        if platform.system() == 'Darwin':  # macOS
            font_paths = [
                '/System/Library/Fonts/PingFang.ttc',
                '/System/Library/Fonts/STHeiti Light.ttc',
                '/Library/Fonts/Arial Unicode.ttf'
            ]
        else:
            font_paths = [
                '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
                '/usr/share/fonts/truetype/arphic/uming.ttc'
            ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont(self.font_name, font_path))
                break

    def convert_all(self):
        """转换目录下所有Word文档为PDF"""
        try:
            input_dir = self.config.output_dir
            for filename in os.listdir(input_dir):
                if filename.endswith('.docx'):
                    word_path = os.path.join(input_dir, filename)
                    pdf_path = os.path.join(input_dir, filename.rsplit('.', 1)[0] + '.pdf')
                    self._convert_to_pdf(word_path, pdf_path)
                    print(f"已生成PDF文件: {pdf_path}")
        except Exception as e:
            print(f"PDF转换过程出错: {str(e)}")

    def _convert_to_pdf(self, word_path, pdf_path):
        """转换单个Word文档为PDF"""
        try:
            doc = Document(word_path)
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=50,
                leftMargin=50,
                topMargin=50,
                bottomMargin=50
            )

            # 创建样式
            styles = {
                'normal': ParagraphStyle(
                    'normal',
                    fontName=self.font_name,
                    fontSize=12,
                    leading=14,
                    alignment=TA_LEFT
                ),
                'center': ParagraphStyle(
                    'center',
                    fontName=self.font_name,
                    fontSize=12,
                    leading=14,
                    alignment=TA_CENTER
                ),
                'red': ParagraphStyle(
                    'red',
                    fontName=self.font_name,
                    fontSize=12,
                    leading=14,
                    alignment=TA_LEFT,
                    textColor=red
                )
            }

            # 处理文档内容
            story = []
            for paragraph in doc.paragraphs:
                # 检查段落对齐方式
                if paragraph.alignment == 1:  # 居中对齐
                    style = styles['center']
                else:
                    style = styles['normal']

                # 处理段落中的红色文本
                text = ''
                for run in paragraph.runs:
                    if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                        # 红色文本
                        text += f'<font color="red">{run.text}</font>'
                    else:
                        text += run.text

                if text.strip():
                    p = Paragraph(text, style)
                    story.append(p)
                    story.append(Spacer(1, 12))  # 段落间距

            # 处理表格（如果有）
            for table in doc.tables:
                # 这里可以添加表格处理逻辑
                pass

            # 生成PDF
            pdf_doc.build(story)

        except Exception as e:
            print(f"转换PDF失败: {word_path} - {str(e)}") 