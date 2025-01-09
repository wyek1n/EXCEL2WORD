from docx import Document
from docx.shared import RGBColor
import os
from .utils import format_filename
from tqdm import tqdm
from .excel_handler import ExcelHandler
import pandas as pd

class WordHandler:
    def __init__(self, config):
        self.config = config
        self.doc = None

    def process_documents(self):
        try:
            # 检查并设置Excel文件路径
            if not self.config.output_dir:
                raise ValueError("输出目录未设置")
            
            excel_path = os.path.join(self.config.output_dir, 'template.xlsx')
            if not os.path.exists(excel_path):
                raise FileNotFoundError(f"Excel文件不存在: {excel_path}")
            
            # 初始化ExcelHandler
            self.config.excel_path = excel_path
            excel_handler = ExcelHandler(self.config)
            
            # 读取Excel数据
            print(f"正在读取Excel文件: {excel_path}")
            df = excel_handler.read_data()
            if df.empty:
                raise ValueError("Excel文件中没有数据")
            
            total_rows = len(df)
            print(f"成功读取数据，共 {total_rows} 行")
            
            # 显示进度条
            for index, row in tqdm(df.iterrows(), total=total_rows, desc="处理进度"):
                try:
                    # 打开一个新的Word文档副本
                    self.doc = Document(self.config.word_template)
                    
                    # 对每个替换项进行处理
                    replace_items = []
                    for item in self.config.replace_items:
                        # 如果项目中包含分号，则分割成多个项目
                        if '；' in item:
                            replace_items.extend([i.strip() for i in item.split('；') if i.strip()])
                        else:
                            replace_items.append(item.strip())
                    
                    # 检查是否所有需要的列都存在
                    missing_columns = [item for item in replace_items if item not in df.columns]
                    if missing_columns:
                        raise ValueError(f"数据中缺少列: {', '.join(missing_columns)}")
                    
                    # 执行替换
                    for item in replace_items:
                        self._replace_text(item, str(row[item]))
                    
                    # 生成输出文件名
                    output_filename = self._generate_filename(row)
                    output_path = os.path.join(self.config.output_dir, output_filename)
                    
                    # 保存文档
                    self.doc.save(output_path)
                    
                except Exception as e:
                    print(f"\n处理第 {index + 1} 行数据时出错: {str(e)}")
                    continue
                    
        except Exception as e:
            raise Exception(f"处理文档失败: {str(e)}")

    def _replace_text(self, old_text, new_text):
        """替换文档中的文本"""
        # 获取字体颜色设置，默认为红色
        font_color = getattr(self.config, 'font_color', 'red')
        rgb_color = RGBColor(255, 0, 0) if font_color == 'red' else RGBColor(0, 0, 0)
        
        # 处理空值或nan
        if pd.isna(new_text) or str(new_text).strip() == '':
            new_text = 'N/A'
        
        for paragraph in self.doc.paragraphs:
            if old_text in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if old_text in item.text:
                        item.text = item.text.replace(old_text, str(new_text))
                        item.font.color.rgb = rgb_color

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, str(new_text))
                                    run.font.color.rgb = rgb_color

    def _generate_filename(self, row):
        """生成输出文件名"""
        filename = self.config.output_format
        
        # 替换文件名中的标记
        replace_items = []
        for item in self.config.replace_items:
            if '；' in item:
                replace_items.extend([i.strip() for i in item.split('；') if i.strip()])
            else:
                replace_items.append(item.strip())
                
        for item in replace_items:
            if item in filename:
                filename = filename.replace(item, str(row[item]))
                
        # 确保文件名有.docx后缀
        if not filename.lower().endswith('.docx'):
            filename += '.docx'
            
        return filename 